"""
Document Generation Module
Génère les documents Word en utilisant les templates existants
Compatible avec l'ancien système generate_files.py
Corrigé :
  ✅ Supprime les espaces vides dans les tableaux
  ✅ Recherche fiable de l'enseignant par ID (code_smartex_ens)
  ✅ Élimine les lignes vides orphelines dans les tableaux
  ✅ Mise en page avec couleurs et formatage
"""

import sys
import json
from docx import Document
from datetime import datetime
import zipfile
import os
from io import BytesIO
import re
import pandas as pd
from docx.shared import RGBColor
from docx2pdf import convert
import tempfile


# ============================================================================
# UTILITAIRE DE CHEMIN POUR PYINSTALLER
# ============================================================================

def get_resource_path(relative_path):
    """
    Obtient le chemin absolu d'une ressource.
    Cherche d'abord à côté de l'exe, puis dans _MEIPASS
    """
    # 1. D'abord, chercher à côté de l'executable
    if getattr(sys, 'frozen', False):
        # Running in PyInstaller bundle
        exe_dir = os.path.dirname(sys.executable)
        external_path = os.path.join(exe_dir, relative_path)

        if os.path.exists(external_path):
            print(f"✅ Template trouvé à côté de l'exe: {external_path}", file=sys.stderr)
            return external_path

        # 2. Sinon, chercher dans _MEIPASS (à l'intérieur de l'exe)
        internal_path = os.path.join(sys._MEIPASS, relative_path)
        if os.path.exists(internal_path):
            print(f"✅ Template trouvé dans l'exe: {internal_path}", file=sys.stderr)
            return internal_path

        print(f"❌ Template introuvable: {relative_path}", file=sys.stderr)
        print(f"   Cherché dans: {exe_dir}", file=sys.stderr)
        print(f"   Cherché dans: {sys._MEIPASS}", file=sys.stderr)

        return internal_path
    else:
        # Running in normal Python
        base_path = os.path.abspath(".")
        return os.path.join(base_path, relative_path)
# ============================================================================
# CHARGEMENT AUTOMATIQUE DU FICHIER DES ENSEIGNANTS
# ============================================================================

def load_enseignants_mapping(excel_dir):
    """
    Charge le fichier 'Enseignants_participants.xlsx'
    depuis le même dossier que le fichier Excel de planning.
    """
    enseignants_file = os.path.join(excel_dir, "Enseignants_participants.xlsx")

    if not os.path.exists(enseignants_file):
        print(json.dumps({
            'success': False,
            'error': f"Fichier enseignants non trouvé: {enseignants_file}"
        }))
        sys.exit(1)

    df = pd.read_excel(enseignants_file)
    enseignants_dict = {}

    for _, row in df.iterrows():
        try:
            id_prof = int(row["code_smartex_ens"])
            nom = str(row["nom_ens"]).strip().title()
            prenom = str(row["prenom_ens"]).strip().title()
            grade = str(row["grade_code_ens"]).strip().upper()
            enseignants_dict[id_prof] = f"{nom} {prenom} ({grade})"
        except Exception:
            continue

    return enseignants_dict

# ============================================================================
# OUTILS DE REMPLACEMENT DE TEXTE
# ============================================================================

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    Remplace du texte dans un paragraphe en préservant le formatage.
    Gère les cas où le texte est fragmenté entre plusieurs runs.
    """
    full_text = paragraph.text
    if old_text.lower() in full_text.lower():
        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
        new_full_text = pattern.sub(new_text, full_text)
        if new_full_text != full_text:
            # Sauvegarder le style du premier run s'il existe
            first_run_font = None
            if paragraph.runs:
                first_run_font = paragraph.runs[0].font

            # Effacer tous les runs
            for run in paragraph.runs:
                run.text = ""

            # Créer un nouveau run avec le texte modifié
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
            else:
                paragraph.add_run(new_full_text)

            return True
    return False


def replace_text_in_document(doc, old_text, new_text):
    """
    Remplace du texte dans tout le document (paragraphes).
    Insensible à la casse.
    """
    replacements = 0
    for paragraph in doc.paragraphs:
        if replace_text_in_paragraph(paragraph, old_text, new_text):
            replacements += 1
    return replacements


def replace_text_in_document_san(doc, old_text):
    """
    Remplace les occurrences de [sance] par S1, S2, S3, etc.
    """
    replacements = 0
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            new_text = 'S' + str(replacements + 1)
            replace_text_in_paragraph(paragraph, old_text, new_text)
            replacements += 1
    return replacements


# ============================================================================
# TRAITEMENT DES DOCUMENTS PAR JOUR (AVEC MISE EN PAGE)
# ============================================================================

def process_day_document(template_path, data, session_type):
    """
    Traite les documents par jour avec correction des lignes vides et mise en page.
    Supprime TOUTES les lignes du modèle avant d'ajouter les données réelles
    Applique la mise en page (couleurs, formatage)
    """
    doc = Document(template_path)
    replace_text_in_document(doc, "[smstre]", "2")
    replace_text_in_document(doc, "[session]", session_type)
    replace_text_in_document(doc, "[annee]", "2024-2025")
    replace_text_in_document(doc, "[date]", datetime.now().strftime("%d/%m/%Y"))
    replace_text_in_document_san(doc, "[sance]")

    table_count = 0
    for table in doc.tables:
        if len(table.columns) == 3:
            table_count += 1
            session_key = f's{table_count}'
            session_data = data.get(session_key, [])
            if not session_data:
                continue

            # ÉTAPE 1: Supprimer TOUTES les lignes vides du modèle
            # (garde seulement la ligne d'en-tête)
            while len(table.rows) > 1:
                tbl = table._tbl
                tbl.remove(tbl.tr_lst[-1])

            # ÉTAPE 2: Ajouter les vraies données avec mise en page
            for teacher_info in session_data:
                row_cells = table.add_row().cells
                if len(teacher_info) >= 2:
                    # Colonne 0: Nom de l'enseignant (texte normal)
                    row_cells[0].text = teacher_info[0]

                    # Colonne 1: Vide
                    row_cells[1].text = ""

                    # Colonne 2: Vide
                    row_cells[2].text = ""

    return doc


# ============================================================================
# TRAITEMENT DES CONVOCATIONS ENSEIGNANTS (AVEC MISE EN PAGE)
# ============================================================================

def process_teacher_document(template_path, prof_name, prof_data):
    """
    Traite les convocations enseignants avec correction des lignes vides et mise en page.
    Supprime TOUTES les lignes du modèle avant d'ajouter les données réelles
    Applique la couleur bleue (RGB 0, 176, 240) aux horaires et durées
    """
    doc = Document(template_path)
    replace_text_in_document(doc, "[prof]", prof_name)

    for table in doc.tables:
        if len(table.columns) == 3:
            # ÉTAPE 1: Supprimer TOUTES les lignes vides du modèle
            # (garde seulement la ligne d'en-tête)
            while len(table.rows) > 1:
                tbl = table._tbl
                tbl.remove(tbl.tr_lst[-1])

            # ÉTAPE 2: Ajouter les vraies données avec mise en page
            for date, surveillances in prof_data.items():
                for surveillance in surveillances:
                    row_cells = table.add_row().cells

                    # Colonne 0: Date (texte normal)
                    row_cells[0].text = date

                    # Colonne 1: Horaires (avec couleur bleue)
                    horaire_text = f"{surveillance[0]} - {surveillance[1]}"
                    # Effacer le texte existant
                    for paragraph in row_cells[1].paragraphs:
                        paragraph.clear()
                    # Ajouter le texte avec couleur
                    run = row_cells[1].paragraphs[0].add_run(horaire_text)
                    run.font.color.rgb = RGBColor(0, 176, 240)

                    # Colonne 2: Durée (avec couleur bleue)
                    duree_text = "1.5h"
                    # Effacer le texte existant
                    for paragraph in row_cells[2].paragraphs:
                        paragraph.clear()
                    # Ajouter le texte avec couleur
                    run = row_cells[2].paragraphs[0].add_run(duree_text)
                    run.font.color.rgb = RGBColor(0, 176, 240)

    # Convert to PDF and return
    # Create temporary files for docx and pdf
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
        temp_docx_path = temp_docx.name
        doc.save(temp_docx_path)

    # Convert to PDF
    temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
    convert(temp_docx_path, temp_pdf_path)

    # Read PDF into BytesIO
    with open(temp_pdf_path, 'rb') as pdf_file:
        pdf_buffer = BytesIO(pdf_file.read())

    # Clean up temporary files
    os.unlink(temp_docx_path)
    os.unlink(temp_pdf_path)

    return pdf_buffer


# ============================================================================
# ORGANISATION DES DONNÉES
# ============================================================================

def organize_data_by_day(planning_data, enseignants_dict):
    """
    Organise les données par jour et par séance.
    """
    days_data = {}
    session_type = "Principale"

    for row in planning_data:
        date = row.get('Date', '')
        time_start = row.get('Heure_Début', '')
        salle = row.get('Salle', '')

        try:
            teacher_id = int(row.get('Enseignant_ID', 0))
        except:
            teacher_id = 0

        teacher_name = enseignants_dict.get(teacher_id, f"Inconnu ({teacher_id})")

        session_key = 's1'
        if '08:30' in str(time_start):
            session_key = 's1'
        elif '10:30' in str(time_start):
            session_key = 's2'
        elif '12:30' in str(time_start):
            session_key = 's3'
        elif '14:30' in str(time_start):
            session_key = 's4'

        if date not in days_data:
            days_data[date] = {'s1': [], 's2': [], 's3': [], 's4': []}

        teacher_info = [teacher_name, salle]
        if teacher_info not in days_data[date][session_key]:
            days_data[date][session_key].append(teacher_info)
            days_data[date][session_key].sort(key=lambda x: x[0])

    return days_data, session_type


def organize_data_by_teacher(planning_data, enseignants_dict):
    """
    Organise les données par enseignant.
    """
    teachers_data = {}

    for row in planning_data:
        try:
            teacher_id = int(row.get('Enseignant_ID', 0))
        except:
            teacher_id = 0

        teacher_name = enseignants_dict.get(teacher_id, f"Inconnu ({teacher_id})")
        teacher_key = f"{teacher_id}::{teacher_name}"

        date = row.get('Date', '')
        time_start = row.get('Heure_Début', '')
        time_fin = row.get('Heure_Fin', '')
        salle = row.get('Salle', '')

        if teacher_key not in teachers_data:
            teachers_data[teacher_key] = {}

        if date not in teachers_data[teacher_key]:
            teachers_data[teacher_key][date] = []

        try:
            debut = datetime.strptime(time_start, "%H:%M")
            fin = datetime.strptime(time_fin, "%H:%M")
            duree = fin - debut
            duree_str = str(duree)
        except:
            duree_str = "1:30:00"

        surveillance_info = [time_start, time_fin, duree_str, salle]
        teachers_data[teacher_key][date].append(surveillance_info)
        teachers_data[teacher_key][date].sort(key=lambda x: x[0])

    for key in teachers_data:
        teachers_data[key] = dict(sorted(
            teachers_data[key].items(),
            key=lambda item: datetime.strptime(item[0], "%d/%m/%Y")
        ))

    return teachers_data


# ============================================================================
# GÉNÉRATION DES DOCUMENTS
# ============================================================================

def generate_global_documents(planning_data, excel_dir, output_dir):
    """
    Génère tous les documents dans un ZIP.
    """
    try:
        enseignants_dict = load_enseignants_mapping(excel_dir)

        # ✅ Debug - Afficher où on cherche
        print(f"🔍 Recherche du template...", file=sys.stderr)
        print(f"🔍 sys._MEIPASS exists: {hasattr(sys, '_MEIPASS')}", file=sys.stderr)
        if hasattr(sys, '_MEIPASS'):
            print(f"🔍 sys._MEIPASS: {sys._MEIPASS}", file=sys.stderr)
        print(f"🔍 __file__: {__file__}", file=sys.stderr)
        print(f"🔍 os.getcwd(): {os.getcwd()}", file=sys.stderr)

        template_source = get_resource_path('enseignansParSeance.docx')
        print(f"🔍 Template source: {template_source}", file=sys.stderr)
        print(f"🔍 Template exists: {os.path.exists(template_source)}", file=sys.stderr)

        # Si le template n'existe pas, lister les fichiers disponibles
        if not os.path.exists(template_source):
            if hasattr(sys, '_MEIPASS'):
                print(f"🔍 Fichiers dans _MEIPASS:", file=sys.stderr)
                for f in os.listdir(sys._MEIPASS):
                    print(f"  - {f}", file=sys.stderr)

            return {'success': False, 'error': f'Template non trouvé: {template_source}'}

        days_data, session_type = organize_data_by_day(planning_data, enseignants_dict)

        semester = "2"
        session = "Principale"
        year = "2024-2025"
        zip_filename = f"affectation_S{semester}_{session}_{year}.zip"
        zip_path = os.path.join(output_dir, zip_filename)

        docs_created = 0
        convocations_created = 0

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for date, data in days_data.items():
                if '/' not in date:
                    continue
                doc = process_day_document(template_source, data, session_type)  # ✅ Utiliser template_source
                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                filename = f"Jour_{date.replace('/', '-')}.docx"
                zipf.writestr(filename, docx_buffer.getvalue())
                docs_created += 1

            teachers_data = organize_data_by_teacher(planning_data, enseignants_dict)

            # ✅ Obtenir le template de convocation
            conv_template_source = get_resource_path('Convocation.docx')

            if os.path.exists(conv_template_source):
                for key, prof_data in teachers_data.items():
                    teacher_id, teacher_name = key.split("::", 1)
                    safe_name = re.sub(r'[^a-zA-Z0-9_]+', '_', teacher_name)
                    conv_filename = f"{safe_name}_S{semester}_{session}_{year}.pdf"
                    pdf_buffer = process_teacher_document(conv_template_source, teacher_name, prof_data)  # ✅ Utiliser conv_template_source
                    pdf_buffer.seek(0)
                    zipf.writestr(conv_filename, pdf_buffer.getvalue())
                    convocations_created += 1

        return {
            'success': True,
            'file': zip_path,
            'days_count': docs_created,
            'convocations_count': convocations_created,
            'message': f'{docs_created} documents journaliers et {convocations_created} convocations générés'
        }

    except Exception as e:
        return {'success': False, 'error': f'Erreur: {str(e)}'}


def generate_teacher_document(planning_data, teacher_id, excel_dir, output_dir):
    """
    Génère le document pour un enseignant spécifique.
    """
    try:
        enseignants_dict = load_enseignants_mapping(excel_dir)
        teacher_id = str(int(teacher_id))

        # ✅ Obtenir le template depuis les ressources PyInstaller
        template_source = get_resource_path('Convocation.docx')

        if not os.path.exists(template_source):
            return {'success': False, 'error': f'Template non trouvé: {template_source}'}

        teachers_data = organize_data_by_teacher(planning_data, enseignants_dict)
        teacher_name = None
        teacher_data = None

        for key, data in teachers_data.items():
            if key.startswith(f"{teacher_id}::"):
                teacher_name = key.split("::", 1)[1]
                teacher_data = data
                break

        if not teacher_data:
            return {'success': False, 'error': f"Aucune surveillance pour l'enseignant ID {teacher_id}"}

        pdf_buffer = process_teacher_document(template_source, teacher_name, teacher_data)  # ✅ Utiliser template_source

        safe_name = re.sub(r'[^a-zA-Z0-9_]+', '_', teacher_name)
        semester = "2"
        session = "Principale"
        year = "2024-2025"
        filename = f"{safe_name}_S{semester}_{session}_{year}.pdf"
        output_path = os.path.join(output_dir, filename)

        with open(output_path, 'wb') as f:
            f.write(pdf_buffer.getvalue())

        surveillances_count = sum(len(v) for v in teacher_data.values())
        return {
            'success': True,
            'file': output_path,
            'teacher_name': teacher_name,
            'surveillances_count': surveillances_count,
            'message': f'Document généré pour {teacher_name}'
        }

    except Exception as e:
        return {'success': False, 'error': str(e)}

# ============================================================================
# 🔹 POINT D'ENTRÉE PRINCIPAL
# ============================================================================

def main():
    if len(sys.argv) < 3:
        print(json.dumps({
            'success': False,
            'error': 'Usage: python generate_docs.py <command> <excel_file> [teacher_id]'
        }))
        return

    command = sys.argv[1]
    data_file = sys.argv[2]

    # ✅ Ajouter cette ligne pour obtenir le dossier du fichier Excel
    excel_dir = os.path.dirname(os.path.abspath(data_file))

    try:
        df = pd.read_excel(data_file)
        planning_data = df.to_dict('records')
    except Exception as e:
        print(json.dumps({'success': False, 'error': f'Erreur de lecture du fichier: {str(e)}'}))
        return

    # Utilise le dossier Téléchargements par défaut
    downloads_dir = os.path.expanduser("~/Downloads")

    # Crée le dossier s'il n'existe pas
    if not os.path.exists(downloads_dir):
        os.makedirs(downloads_dir)

    output_dir = downloads_dir

    if command == 'global':
        result = generate_global_documents(planning_data, excel_dir, output_dir)  # ✅ Ajouter excel_dir
    elif command == 'teacher':
        if len(sys.argv) < 4:
            print(json.dumps({'success': False, 'error': 'ID enseignant manquant'}))
            return
        teacher_id = sys.argv[3]
        result = generate_teacher_document(planning_data, teacher_id, excel_dir, output_dir)  # ✅ Ajouter excel_dir
    else:
        result = {'success': False, 'error': f'Commande inconnue: {command}'}

    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()